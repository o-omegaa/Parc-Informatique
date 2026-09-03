from __future__ import annotations

import re
from pathlib import Path

import pypdfium2 as pdfium
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Users\lenovo\Desktop\supplier-portal")
SOURCE = Path(r"C:\Users\lenovo\Desktop\rapport\rapport_stage_dgssi.md")
REFERENCE_PDF = Path(r"C:\Users\lenovo\Downloads\rapport_stage_dgssi_final (2).pdf")
REFERENCE_PREVIEW_DIR = ROOT / "tmp" / "docs" / "revised-pdf-preview"
OUTPUT = ROOT / "output" / "docx" / "rapport_stage_dgssi_final_version_word.docx"
ASSET_DIR = ROOT / "tmp" / "docs" / "docx-assets-revised"

NAVY = "000000"
NAVY_LIGHT = "000000"
GREEN = "009540"
TEXT = "000000"
MID = "000000"
PALE = "F4F6F9"
WHITE = "FFFFFF"
A4_WIDTH_CM = 21
A4_HEIGHT_CM = 29.7
MARGIN_X_CM = 2.05
MARGIN_TOP_CM = 1.75
MARGIN_BOTTOM_CM = 1.75
CONTENT_WIDTH_CM = A4_WIDTH_CM - (MARGIN_X_CM * 2)
CONTENT_WIDTH_DXA = round(CONTENT_WIDTH_CM / 2.54 * 1440)


def set_run_font(run, *, name="Times New Roman", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def ensure_style(doc, name, style_type=WD_STYLE_TYPE.PARAGRAPH, base=None):
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, style_type)
        if base:
            style.base_style = doc.styles[base]
        return style


def configure_style(style, *, font="Times New Roman", size=10.8, color=TEXT,
                    bold=False, italic=False, before=0, after=6, line=1.28,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, keep_with_next=False):
    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:ascii"), font)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor.from_string(color)
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.alignment = alignment
    fmt.keep_with_next = keep_with_next


def set_paragraph_border(paragraph, *, top=None, bottom=None):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    for edge, spec in (("top", top), ("bottom", bottom)):
        if not spec:
            continue
        element = pbdr.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            pbdr.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(spec.get("size", 6)))
        element.set(qn("w:space"), str(spec.get("space", 1)))
        element.set(qn("w:color"), spec.get("color", "000000"))


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=75, start=100, bottom=75, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, **borders):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, spec in borders.items():
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        if spec is None:
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), spec.get("val", "single"))
            element.set(qn("w:sz"), str(spec.get("size", 4)))
            element.set(qn("w:color"), spec.get("color", "000000"))


def set_table_geometry(table, widths_dxa, *, indent=0):
    total = int(sum(widths_dxa))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths_dxa):
        col.set(qn("w:w"), str(int(width)))
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width)))
            tc_w.set(qn("w:type"), "dxa")


def repeat_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def add_field(paragraph, instruction, fallback=""):
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.extend([begin, instr, separate, text, end])
    return run


def set_update_fields_on_open(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    node = sect_pr.find(qn("w:pgNumType"))
    if node is None:
        node = OxmlElement("w:pgNumType")
        sect_pr.append(node)
    node.set(qn("w:start"), str(start))


def set_section_geometry(section):
    section.page_width = Cm(A4_WIDTH_CM)
    section.page_height = Cm(A4_HEIGHT_CM)
    section.left_margin = Cm(MARGIN_X_CM)
    section.right_margin = Cm(MARGIN_X_CM)
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def set_cover_geometry(section):
    """Keep the user-approved PDF cover visually intact as a full-page image."""
    section.page_width = Cm(A4_WIDTH_CM)
    section.page_height = Cm(A4_HEIGHT_CM)
    section.left_margin = Cm(0.15)
    section.right_margin = Cm(0.15)
    section.top_margin = Cm(0.15)
    section.bottom_margin = Cm(0.15)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)


def add_markdown_runs(paragraph, text, *, default_size=None, default_color=TEXT):
    token_pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*|\[[^\]]+\]\([^)]+\))")
    position = 0
    for match in token_pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=default_size, color=default_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, color=default_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Courier New", size=(default_size or 9.0), color=default_color)
        elif token.startswith("["):
            label = re.sub(r"^\[([^\]]+)\]\([^)]+\)$", r"\1", token)
            run = paragraph.add_run(label)
            set_run_font(run, size=default_size, color=default_color, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=default_size, color=default_color, italic=True)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=default_size, color=default_color)


def clean_markdown(text):
    return text.replace("\\*", "*").strip()


def add_rule(doc, *, color="000000", size=6, before=0, after=8, keep_with_next=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    set_paragraph_border(paragraph, bottom={"color": color, "size": size, "space": 1})
    return paragraph


def add_paragraph(doc, text, style="Body", *, alignment=None):
    paragraph = doc.add_paragraph(style=style)
    if alignment is not None:
        paragraph.alignment = alignment
    add_markdown_runs(paragraph, clean_markdown(text))
    return paragraph


def add_picture(doc, path, width_cm, *, caption=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        caption_p = doc.add_paragraph(style="Caption")
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_p.add_run(caption)
        set_run_font(run, size=8.8, color=TEXT)
    return paragraph


def create_diagram_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    source = pdfium.PdfDocument(REFERENCE_PDF) if REFERENCE_PDF.exists() else None
    if source is None and not (REFERENCE_PREVIEW_DIR / "page-01.png").exists():
        raise FileNotFoundError("Ni le PDF révisé ni ses aperçus rendus ne sont disponibles.")
    # pages/crops are expressed as ratios, so they remain stable at high-resolution rendering.
    specs = {
        "cover": (1, (0.0, 0.0, 1.0, 1.0)),
        "architecture": (15, (0.12, 0.35, 0.88, 0.525)),
        "clean": (16, (0.17, 0.09, 0.83, 0.32)),
        "usecases": (17, (0.10, 0.09, 0.90, 0.455)),
        "data": (18, (0.10, 0.275, 0.90, 0.695)),
        "security": (24, (0.12, 0.14, 0.88, 0.275)),
        "login": (27, (0.09, 0.24, 0.91, 0.60)),
        "admin": (28, (0.10, 0.085, 0.90, 0.42)),
        "purchase": (28, (0.10, 0.60, 0.90, 0.90)),
        "supplier": (29, (0.10, 0.22, 0.90, 0.47)),
    }
    assets = {}
    for name, (page_number, ratios) in specs.items():
        output = ASSET_DIR / f"{name}.png"
        if source is not None:
            page = source[page_number - 1]
            image = page.render(scale=3.2).to_pil()
        else:
            image = Image.open(REFERENCE_PREVIEW_DIR / f"page-{page_number:02d}.png").convert("RGB")
        width, height = image.size
        x1, y1, x2, y2 = ratios
        crop = image.crop((int(width * x1), int(height * y1), int(width * x2), int(height * y2)))
        crop.save(output)
        assets[name] = output
    return assets


def build_styles(doc):
    styles = doc.styles
    configure_style(styles["Normal"], font="Times New Roman", size=10.8, color=TEXT,
                    before=0, after=6, line=1.28, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    styles["Normal"].paragraph_format.widow_control = True

    heading1 = styles["Heading 1"]
    configure_style(heading1, font="Times New Roman", size=27, color=NAVY, bold=True,
                    before=9, after=13, line=1.06, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    keep_with_next=True)
    heading1.paragraph_format.keep_together = True
    heading2 = styles["Heading 2"]
    configure_style(heading2, font="Times New Roman", size=17, color=NAVY, bold=True,
                    before=16, after=7, line=1.1, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    keep_with_next=True)
    heading2.paragraph_format.keep_together = True
    heading3 = styles["Heading 3"]
    configure_style(heading3, font="Times New Roman", size=13.2, color=NAVY, bold=True,
                    before=12, after=5, line=1.12, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    keep_with_next=True)
    heading3.paragraph_format.keep_together = True

    body = ensure_style(doc, "Body", base="Normal")
    configure_style(body, font="Times New Roman", size=10.8, color=TEXT,
                    before=0, after=6, line=1.28, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    body.paragraph_format.widow_control = True

    chapter = ensure_style(doc, "ChapterLabel", base="Normal")
    configure_style(chapter, font="Times New Roman", size=17, color=NAVY, bold=True,
                    before=0, after=8, line=1.1, alignment=WD_ALIGN_PARAGRAPH.LEFT, keep_with_next=True)

    caption = ensure_style(doc, "Caption", base="Normal")
    configure_style(caption, font="Times New Roman", size=8.8, color=TEXT,
                    before=3, after=8, line=1.08, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    code = ensure_style(doc, "CodeBlock", base="Normal")
    configure_style(code, font="Courier New", size=7.05, color="333333",
                    before=4, after=8, line=1.0, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    code.paragraph_format.left_indent = Cm(0.25)
    code.paragraph_format.right_indent = Cm(0.25)

    toc1 = ensure_style(doc, "TOC1", base="Normal")
    configure_style(toc1, font="Times New Roman", size=10.5, color=NAVY, bold=True,
                    before=0, after=3, line=1.1, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    toc2 = ensure_style(doc, "TOC2", base="Normal")
    configure_style(toc2, font="Times New Roman", size=9.6, color=NAVY,
                    before=0, after=2, line=1.08, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    toc2.paragraph_format.left_indent = Cm(0.55)
    toc3 = ensure_style(doc, "TOC3", base="Normal")
    configure_style(toc3, font="Times New Roman", size=8.8, color=MID,
                    before=0, after=1, line=1.06, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    toc3.paragraph_format.left_indent = Cm(1.1)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        configure_style(style, font="Times New Roman", size=10.6, color=TEXT,
                        before=0, after=4, line=1.23, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def build_body_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(CONTENT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)
    run = paragraph.add_run("Portail Fournisseur")
    set_run_font(run, size=8.2, color=TEXT, italic=True)
    tab_run = paragraph.add_run("\t")
    set_run_font(tab_run, size=8.2, color=TEXT, italic=True)
    run = paragraph.add_run("Rapport de projet de fin d'année")
    set_run_font(run, size=8.2, color=TEXT, italic=True)
    set_paragraph_border(paragraph, bottom={"color": "4E4E4E", "size": 4, "space": 2})

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(0)
    set_paragraph_border(paragraph, top={"color": "4E4E4E", "size": 4, "space": 2})
    run = add_field(paragraph, " PAGE ", "1")
    set_run_font(run, size=9, color=TEXT)


def cover_paragraph(doc, text="", *, size=12, color=TEXT, bold=False, italic=False,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=1.1):
    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)


def create_cover(doc, cover_image):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(1)
    paragraph.add_run().add_picture(str(cover_image), width=Cm(20.7), height=Cm(29.2))


def add_toc_entry(doc, label, page, level=1):
    style = {1: "TOC1", 2: "TOC2", 3: "TOC3"}[level]
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(CONTENT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = paragraph.add_run(label)
    set_run_font(run, size={1: 10.5, 2: 9.6, 3: 8.8}[level], color={1: NAVY, 2: NAVY, 3: MID}[level], bold=(level == 1))
    paragraph.add_run("\t")
    run = paragraph.add_run(str(page))
    set_run_font(run, size={1: 10.5, 2: 9.6, 3: 8.8}[level], color={1: NAVY, 2: NAVY, 3: MID}[level], bold=(level == 1))


def add_static_toc(doc):
    doc.add_paragraph("Table des matières", style="Heading 1")
    entries = [
        ("Introduction générale", 8, 1), ("Contexte général", 8, 2), ("Problématique", 8, 2),
        ("Objectifs du projet", 8, 2), ("Méthodologie de travail", 9, 2), ("Organisation du rapport", 9, 2),
        ("Généralités sur le projet", 10, 1), ("1.1 Introduction", 10, 2),
        ("1.2 Présentation de l'organisme d'accueil", 10, 2), ("1.2.1 Missions principales", 10, 3),
        ("1.2.2 Contexte numérique et enjeux de cybersécurité", 10, 3), ("1.3 Étude de l'existant", 10, 2),
        ("1.4 Expression du besoin", 11, 2), ("1.5 Conclusion", 13, 2),
        ("Conception du projet", 14, 1), ("2.1 Introduction", 14, 2),
        ("2.2 Architecture générale de la solution", 14, 2), ("2.2.1 Architecture en couches du back-end : Clean Architecture", 15, 3),
        ("2.3 Diagramme de cas d'utilisation", 16, 2), ("2.4 Diagramme de classes", 17, 2),
        ("2.5 Conception de la base de données", 18, 2), ("2.6 Conclusion", 19, 2),
        ("Réalisation du projet", 20, 1), ("3.1 Introduction", 20, 2),
        ("3.2 Outils et langages de programmation utilisés", 20, 2), ("3.3 Architecture technique détaillée", 21, 2),
        ("3.4 Sécurité de l'application", 23, 2), ("3.5 Le module de messagerie électronique", 25, 2),
        ("3.6 Interfaces de l'application", 26, 2), ("3.7 Conclusion", 27, 2),
        ("Conclusion générale et perspectives", 28, 1), ("Bibliographie et webographie", 31, 1), ("Annexes", 32, 1),
    ]
    for entry in entries:
        add_toc_entry(doc, *entry)


def add_named_list(doc, title, items):
    doc.add_paragraph(title, style="Heading 1")
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        add_markdown_runs(paragraph, item)


def add_markdown_table(doc, rows):
    rows = [[clean_markdown(cell) for cell in row] for row in rows]
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    rows = [row + [""] * (col_count - len(row)) for row in rows]
    longest = [max(5, max(len(re.sub(r"\*|`", "", row[index])) for row in rows)) for index in range(col_count)]
    total_weight = sum(longest)
    widths = [max(900, round(CONTENT_WIDTH_DXA * value / total_weight)) for value in longest]
    if col_count == 2:
        widths[0] = max(widths[0], 1900)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    table = doc.add_table(rows=len(rows), cols=col_count)
    set_table_geometry(table, widths, indent=0)
    for row_index, source_row in enumerate(rows):
        row = table.rows[row_index]
        prevent_row_split(row)
        if row_index == 0:
            repeat_header_row(row)
        for col_index, value in enumerate(source_row):
            cell = row.cells[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=80, start=100, bottom=80, end=100)
            cell_p = cell.paragraphs[0]
            cell_p.paragraph_format.space_before = Pt(0)
            cell_p.paragraph_format.space_after = Pt(0)
            cell_p.paragraph_format.line_spacing = 1.08
            cell_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_markdown_runs(cell_p, value, default_size=(8.4 if col_count >= 3 else 8.9))
            if row_index == 0:
                shade_cell(cell, PALE)
                for run in cell_p.runs:
                    set_run_font(run, size=(8.3 if col_count >= 3 else 8.8), color=TEXT, bold=True)
                set_cell_borders(cell, top={"color": "000000", "size": 6}, bottom={"color": "000000", "size": 6}, left=None, right=None)
            elif row_index == len(rows) - 1:
                set_cell_borders(cell, bottom={"color": "000000", "size": 6}, left=None, right=None)
            else:
                set_cell_borders(cell, left=None, right=None)
    end = doc.add_paragraph()
    end.paragraph_format.space_after = Pt(1)


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph(style="CodeBlock")
    paragraph.paragraph_format.keep_together = True
    shade_paragraph(paragraph, "F3F5F8")
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, name="Courier New", size=7.05, color="333333")


def parse_report(doc, markdown, assets):
    start = markdown.find("# Remerciements")
    if start < 0:
        raise ValueError("La section 'Remerciements' est introuvable dans le fichier Markdown.")
    lines = markdown[start:].splitlines()
    index = 0
    body = []
    toc_done = False

    diagrams = {
        "2.2 Architecture générale de la solution": ("architecture", "Figure 2.1 : Architecture générale du Portail Fournisseur."),
        "2.2.1 Architecture en couches du back-end : Clean Architecture": ("clean", "Figure 2.2 : Organisation du back-end selon la Clean Architecture."),
        "2.3 Diagramme de cas d'utilisation": ("usecases", "Figure 2.3 : Principaux cas d'utilisation selon les trois profils du système."),
        "2.4 Diagramme de classes": ("data", "Figure 2.4 : Vue synthétique des entités métier principales."),
        "3.4 Sécurité de l'application": ("security", "Figure 3.1 : Parcours sécurisé d'authentification : contrôle des identifiants, JWT, RBAC et journalisation."),
    }

    def flush_body():
        nonlocal body
        if body:
            paragraph_text = " ".join(part.strip() for part in body).strip()
            if paragraph_text:
                add_paragraph(doc, paragraph_text)
            body = []

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped == "\\newpage":
            flush_body()
            next_content = next((line.strip() for line in lines[index + 1:] if line.strip()), "")
            # The added application screenshots already fill the preceding page;
            # a forced break here would create an empty page before chapter 4.
            if next_content != "# Chapitre 4 : Conclusion générale et perspectives":
                doc.add_page_break()
            index += 1
            continue
        if not stripped or stripped == "---":
            flush_body()
            index += 1
            continue
        if stripped.startswith("```"):
            flush_body()
            code = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index].rstrip())
                index += 1
            add_code_block(doc, code)
            index += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            flush_body()
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
                rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            add_markdown_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_body()
            level = len(heading.group(1))
            title = clean_markdown(heading.group(2))
            if level == 1:
                if title == "Liste des acronymes" and not toc_done:
                    add_static_toc(doc)
                    doc.add_page_break()
                    add_named_list(doc, "Liste des figures", [
                        "Figure 2.1 : Architecture générale du Portail Fournisseur",
                        "Figure 2.2 : Organisation du back-end selon la Clean Architecture",
                        "Figure 2.3 : Cas d'utilisation selon les profils utilisateur",
                        "Figure 2.4 : Vue synthétique des entités métier",
                        "Figure 3.1 : Parcours sécurisé d'authentification : contrôle, JWT, RBAC et audit",
                        "Figure 3.2 : Interfaces du Portail Fournisseur selon les profils utilisateur",
                    ])
                    doc.add_page_break()
                    add_named_list(doc, "Liste des tableaux", [
                        "Tableau 1.1 : Acteurs et responsabilités principales",
                        "Tableau 1.2 : Synthèse des besoins non fonctionnels",
                        "Tableau 2.1 : Description des classes principales",
                        "Tableau 3.1 : Technologies back-end",
                        "Tableau 3.2 : Technologies front-end",
                        "Tableau 3.3 : Principaux points d'entrée de l'API REST",
                        "Tableau 3.4 : Matrice de contrôle d'accès par rôle",
                    ])
                    doc.add_page_break()
                    toc_done = True
                chapter_match = re.match(r"Chapitre\s+(\d+)\s*:\s*(.+)", title, re.IGNORECASE)
                if chapter_match:
                    label = doc.add_paragraph(style="ChapterLabel")
                    run = label.add_run(f"CHAPITRE {chapter_match.group(1)}")
                    set_run_font(run, size=17, color=NAVY, bold=True)
                    add_rule(doc, color=NAVY_LIGHT, size=6, after=11, keep_with_next=True)
                    title = chapter_match.group(2)
                elif title in {"Introduction générale", "Bibliographie et webographie", "Annexes"}:
                    add_rule(doc, color=NAVY_LIGHT, size=6, after=11)
                paragraph = doc.add_paragraph(style="Heading 1")
                add_markdown_runs(paragraph, title, default_size=27, default_color=NAVY)
            elif level == 2:
                paragraph = doc.add_paragraph(style="Heading 2")
                add_markdown_runs(paragraph, title, default_size=17, default_color=NAVY)
            else:
                paragraph = doc.add_paragraph(style="Heading 3")
                add_markdown_runs(paragraph, title, default_size=13.2, default_color=NAVY)

            if title in diagrams:
                key, caption = diagrams[title]
                add_picture(doc, assets[key], 15.0, caption=caption)
            if title == "3.6.1 Interface de connexion":
                add_picture(doc, assets["login"], 15.0, caption="Figure 3.2 : Aperçu de l'écran de connexion du Portail Fournisseur.")
            elif title == "3.6.2 Interface d'administration":
                add_picture(doc, assets["admin"], 15.0, caption="Aperçu du tableau de bord administrateur.")
            elif title == "3.6.3 Interface du service achat":
                add_picture(doc, assets["purchase"], 15.0, caption="Aperçu de l'espace du service achat.")
            elif title == "3.6.4 Interface fournisseur":
                add_picture(doc, assets["supplier"], 15.0, caption="Aperçu de l'espace fournisseur.")
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        ordered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if bullet or ordered:
            flush_body()
            ordered_list = bool(ordered)
            items = []
            while index < len(lines):
                candidate = lines[index].strip()
                match = re.match(r"^(\d+)\.\s+(.+)$", candidate) if ordered_list else re.match(r"^[-*]\s+(.+)$", candidate)
                if not match:
                    break
                items.append(match.group(2) if ordered_list else match.group(1))
                index += 1
            for item in items:
                paragraph = doc.add_paragraph(style="List Number" if ordered_list else "List Bullet")
                add_markdown_runs(paragraph, clean_markdown(item))
            continue
        if stripped.startswith("*") and stripped.endswith("*") and len(stripped) > 2:
            flush_body()
            is_bold = stripped.startswith("**") and stripped.endswith("**")
            text = stripped[2:-2].strip() if is_bold else stripped[1:-1].strip()
            if text.startswith("Tableau"):
                paragraph = doc.add_paragraph(style="Caption")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(text)
                set_run_font(run, size=8.8, color=TEXT)
            else:
                paragraph = doc.add_paragraph(style="Body")
                run = paragraph.add_run(text)
                set_run_font(run, size=10.8, color=TEXT, bold=is_bold, italic=not is_bold)
            index += 1
            continue
        body.append(raw)
        index += 1
    flush_body()


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(f"Fichier Markdown introuvable : {SOURCE}")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_update_fields_on_open(doc)
    doc.core_properties.title = "Rapport de stage - Portail Fournisseur"
    doc.core_properties.author = "Hamza AIT OUMGHAR"
    doc.core_properties.subject = "Rapport de Projet de Fin d'Année"
    doc.core_properties.comments = "Version Word du rapport final."
    build_styles(doc)
    assets = create_diagram_assets()
    cover_section = doc.sections[0]
    set_cover_geometry(cover_section)
    cover_section.different_first_page_header_footer = True
    create_cover(doc, assets["cover"])

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(body_section)
    body_section.different_first_page_header_footer = False
    set_page_number_start(body_section, 1)
    build_body_header_footer(body_section)

    parse_report(doc, SOURCE.read_text(encoding="utf-8"), assets)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
